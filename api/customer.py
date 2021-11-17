from api import EMPTY_RESPONSE
from core import flask_api, db
import flask
import flask_restful
import models
import schema
from datetime import date, timedelta, datetime
import errors
from utils import general, decorators
from docx import Document
import time
from flask import current_app


@flask_api.resource('/customer', '/customer/<int:customer_id>')
class CustomerResource(flask_restful.Resource):
    @staticmethod
    @decorators.check_session_role()
    def get(customer_id=None):
        """
        Get customer by id or all customers
        :param customer_id:
        :return:
        """
        validated_data = schema.CustomerEditRequestSchema().load(flask.request.args or {})

        customer = models.Customer.get_by_filters(
            customer_id=validated_data.get('id'),
            customer_type=validated_data.get('type')
        )

        if customer_id:
            return schema.CustomerSchema(many=False).dump(models.Customer.get_by_id(customer_id))

        return schema.CustomerSchema(many=True).dump(customer)

    @staticmethod
    @decorators.check_session_role()
    def post():
        """
        Create customer by given data
        :return:
        """
        validated_data = schema.CustomerRequestSchema().load(flask.request.json or {})

        # Create customer
        customer = models.Customer.create(**validated_data)

        return schema.CustomerSchema(many=False).dump(customer)

    @staticmethod
    @decorators.check_session_role()
    def patch(customer_id):
        """
        Edit customer for given customer id
        :param customer_id:
        :return:
        """
        validated_data = schema.CustomerEditRequestSchema().load(flask.request.json or {})

        customer = general.get_object_by_id(
            obj_id=customer_id,
            query_model=models.Customer,
            error=errors.ERR_BAD_CUSTOMER_ID
        )

        # Edit customer
        customer.edit(**validated_data)

        return schema.CustomerSchema(many=False).dump(customer)

    @staticmethod
    @decorators.check_session_role(models.UserRoleEnum.admin, check_role=True)
    def delete(customer_id):
        """
        Delete customer for given customer id
        """
        customer = general.get_object_by_id(
                obj_id=customer_id,
                query_model=models.Customer,
                error=errors.ERR_BAD_CUSTOMER_ID
            )

        # Delete customer
        customer.edit(deleted=True)

        return EMPTY_RESPONSE


@flask_api.resource('/customer/<int:customer_id>/apartment/<int:apartment_id>/contract',
                    '/customer/<int:customer_id>/apartment/<int:apartment_id>/contract/<int:contract_id>',)
class ContractResource(flask_restful.Resource):
    @staticmethod
    @decorators.check_session_role()
    def get(customer_id, apartment_id, contract_id):
        """
        Get contract by given data
        :params contract_id, apartment_id, customer_id:
        :return:
        """

        return schema.ContractSchema(many=False).dump(
            models.Contract.get_by_customer_and_apartment_with_contract(customer_id=customer_id,
                                                                        contract_id=contract_id,
                                                                        apartment_id=apartment_id)
            )

    @staticmethod
    @decorators.check_session_role(return_user=True)
    def post(current_user, customer_id, apartment_id):
        """
        Create new contract by given data
        :params current_user, customer_id, apartment_id:
        :return:
        """
        validated_data = schema.ContractOptionalRequestSchema().load(flask.request.json or {})

        validated_data['user_id'] = current_user.id
        validated_data['customer_id'] = customer_id
        validated_data['apartment_id'] = apartment_id

        apartment = general.get_object_by_id(
            obj_id=apartment_id,
            query_model=models.Apartment,
            error=errors.ERR_BAD_APARTMENT_ID
        )

        contract = models.Contract.get_by_customer_and_apartment(customer_id=customer_id, apartment_id=apartment_id)

        if contract:
            flask_restful.abort(400, error=errors.ERR_CONTRACT_ALREADY_EXISTS)

        # Get sold apartments by contracts
        contract_issued = models.Contract.get_issued_by_apartment(
            apartment_id=apartment_id
        )

        if contract_issued:
            flask_restful.abort(400, error=errors.ERR_APARTMENT_NOT_AVAILABLE)

        if apartment.status.name in ['sold', 'reserved']:
            flask_restful.abort(400, error=errors.ERR_APARTMENT_NOT_AVAILABLE)

        validated_data['price'] = apartment.price
        validated_data['contract_number'] = '{}'.format(str(time.time()).split('.')[0])

        if validated_data.get('approved') is True and current_user.role.name != 'finance':
            flask_restful.abort(400, error=errors.ERR_BAD_ROLE_FOR_APPROVING)

        contract = models.Contract.create(**validated_data)

        return schema.ContractSchema(many=False).dump(contract)

    @staticmethod
    @decorators.check_session_role(return_user=True)
    def patch(current_user, customer_id, apartment_id, contract_id):
        """
        Edit contract for given contract id
        :param current_user:
        :param customer_id:
        :param apartment_id:
        :param contract_id:
        :return:
        """
        validated_data = schema.ContractOptionalRequestSchema().load(flask.request.json or {})

        validated_data['user_id'] = current_user.id
        validated_data['customer_id'] = customer_id
        validated_data['apartment_id'] = apartment_id

        contract = models.Contract.get_by_contract_number(contract_number=validated_data.get('contract_number'))

        if contract:
            flask_restful.abort(400, error=errors.ERR_CONTRACT_ALREADY_EXISTS)

        contract = general.get_object_by_id(
            obj_id=contract_id,
            query_model=models.Contract,
            error=errors.ERR_BAD_CONTRACT_ID
        )

        contract.edit(price=validated_data.get('price'),
                      status=validated_data.get('status'),
                      note=validated_data.get('note'))

        if current_user.role.name == 'finance' and validated_data.get('approved'):
            contract.edit(approved_by=current_user.id, approved=True)
        elif current_user.role.name != 'finance' and validated_data.get('approved'):
            flask_restful.abort(400, error=errors.ERR_BAD_ROLE_FOR_APPROVING)

        if contract.apartment.price != contract.price and not contract.approved:
            flask_restful.abort(400, error=errors.ERR_PRICE_NOT_APPROVED)

        if contract.status == 'reserved' or validated_data.get('status') == 'reserved':
            contract.apartment.edit(status='reserved')

            # Get all non customer contracts by apartment
            contracts = models.Contract.get_all_non_customer_by_apartment(
                apartment_id=apartment_id,
                customer_id=customer_id
            )

            # Delete all non customer apartment contracts
            for agreement in contracts:
                agreement.edit(deleted=True)

            # Create word template
            file_name = 'Ugovor o kupoprodaji br. {}{}'.format(contract.id, '.docx')
            path = '{}{}'.format(current_app.config.get('CONTRACT_PATH'), file_name)
            document = Document()
            # section = document.sections[0]
            # header = section.header
            # p = header.paragraphs[0]
            # p.text = '\tUGOVOR O KUPOPRODAJI NEPOKRETNOSTI\t'
            document.add_heading('\tUGOVOR O KUPOPRODAJI NEPOKRETNOSTI\t', level=1)
            p = document.add_paragraph(' zaključen dana {} godine, između ugovornih strana:'
                                       .format(contract.date_of_update.strftime('%d.%m.%Y'))
                                       )
            # Adding user salesman credentials
            p = document.add_paragraph('1.{} kao prodavca'.format(f'{contract.user.first_name} {contract.user.last_name}'))

            # Adding customer credentials
            p = document.add_paragraph('2.{} , ul. {}, kao kupca, sa druge strane'.format(contract.customer.name, contract.customer.address))

            # p.text = '\tUGOVORNE STRANE SU SE SPORAZUMELE O SLEDEĆEM:\t'
            document.add_heading('\tUGOVORNE STRANE SU SE SPORAZUMELE O SLEDEĆEM:\t', level=3)

            p = document.add_paragraph('Prodavac prodaje, a kupac kupuje stan br. {} po medjusobno ugovorenoj ceni od {} $.'
                                       .format(contract.apartment.id, contract.price))

            if validated_data.get('signed'):
                contract.apartment.edit(status='sold')
                contract.edit(status='purchased', signed=True)

                # Add contract details
                p.add_run(' Način placanja: {}, na osnovu ugovora pod brojem {}'.format(contract.payment_method.name, contract.contract_number))

            document.save(path)

        return schema.ContractSchema(many=False).dump(contract)

    @staticmethod
    @decorators.check_session_role(models.UserRoleEnum.finance, check_role=True)
    def delete(customer_id, apartment_id, contract_id):
        """
        Delete contract for given contract id
        """
        contract = general.get_object_by_id(
            obj_id=contract_id,
            query_model=models.Contract,
            error=errors.ERR_BAD_CONTRACT_ID
        )

        # Delete contract
        contract.edit(deleted=True)

        return EMPTY_RESPONSE
